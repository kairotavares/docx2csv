#!/usr/bin/env python
# -*- coding: utf8 -*-

import click

import os
import csv
import sys
import xlwt

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def extract_table(table):
    """Extracts table data from table object"""
    results = []
    for row in table.rows:
        r = []
        for cell in row.cells:
            r.append(cell.text.replace(u'\n', u' ').encode('utf8'))
        results.append(r)
    return results


def extract_docx_table(filename):
    """Extracts table from .DOCX files"""
    tables = []
    document = Document(filename)
#    print dir(document)
#    print document.tables
    n = 0
    for table in document.tables:
        n += 1
#        print '## TABLE %d ##' % (n)
        tdata = extract_table(table)
        tables.append(tdata)
    return tables

def store_table(tabdata, filename, format='csv'):
    """Saves table data as csv file"""
    if format == 'csv':
        f = file(filename, 'w')
        w = csv.writer(f, delimiter=',')
        for row in tabdata:
            w.writerow(row)
    elif format == 'xls':
        workbook = xlwt.Workbook()
        sheet = workbook.add_sheet('0')
        rn = 0
        for row in tabdata:
            cn = 0
            for c in row:
                sheet.write(rn, cn, c.decode('utf8'))
                cn += 1
            rn += 1
        workbook.save(filename)


@click.group()
def cli1():
    """Extracts tables from DOCX files as CSV or XLSX.
        Use command: "docx2csv convert <filename>" to run extraction.
        It will create files like filename_1.csv, filename_2.csv for each table found.

    """
    pass

@cli1.command()
@click.option('--format', default='csv', help='Output format: CSV, XLSX')
@click.option('--singlefile', default=False, help='Outputs XLS file with multiple sheets' )
@click.option('--sizefilter', default=0, help='Filters table by size number of rows')
@click.argument('filename')
def extract(format, sizefilter, singlefile, filename):
    """Extracts tables from DOCX files as CSV or XLSX.

        Use command: "docx2csv convert <filename>" to run extraction.
        It will create files like filename_1.csv, filename_2.csv for each table found.
    """
    tables = extract_docx_table(filename)
    name = filename.rsplit('.', 1)[0]
    format = format.lower()
    n = 0
    lfilter = int(sizefilter)
    if singlefile:
        workbook = xlwt.Workbook()
        for t in tables:
            if lfilter >= len(t):
                print 'Table length %d instead of %d. Skipped' % (len(t), lfilter)
                continue
            n += 1
            sheet = workbook.add_sheet(str(n))
            rn = 0
            for row in t:
                cn = 0
                for c in row:
                    sheet.write(rn, cn, c.decode('utf8'))
                    cn += 1
                rn += 1
        destname = name + '.%s' % (format)
        workbook.save(destname)
        print destname, 'saved'
    else:
        for t in tables:
            if lfilter >= len(t):
                print 'Table length %d instead of %d. Skipped' % (len(t), lfilter)
                continue
            n += 1
            destname = name + '_%d.%s' % (n, format)
            store_table(t, destname, format)
            print destname, 'saved'


#cli = click.CommandCollection(sources=[cli1])

if __name__ == '__main__':
    extract()
